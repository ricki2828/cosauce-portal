"""
Contract Service - Business logic for contract generation
Ported from contract-generator with Anthropic integration
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Dict, List, Optional
from datetime import date, datetime
import json
import anthropic

from ..config import ANTHROPIC_API_KEY, CONTRACT_TEMPLATES_DIR, OUTPUTS_DIR


@dataclass
class ContractData:
    """Data structure for contract form inputs."""
    contract_type: str  # "MSA" | "SOW" | "ShortForm"

    # Common fields
    client_name: str = ""
    client_legal_name: str = ""
    client_address: str = ""
    client_contact_name: str = ""
    client_contact_email: str = ""
    client_contact_phone: str = ""
    effective_date: str = ""

    # Provider details (CoSauce defaults)
    provider_name: str = "Ricki Taiaroa"
    provider_email: str = "ricki@cosauce.co"
    provider_phone: str = "+65 8815 0365"
    provider_company: str = "CoSauce Pte Ltd"
    provider_address: str = "Singapore"

    # MSA-specific
    term_duration: str = "12 months"
    notice_period: str = "60 days"
    jurisdiction: str = "Singapore"
    payment_terms: str = "Net 30"
    currency: str = "SGD"
    custom_clauses: str = ""

    # SOW-specific
    project_name: str = ""
    project_start: str = ""
    project_end: str = ""
    scope_bullets: List[str] = field(default_factory=list)
    scope_formal: str = ""  # AI-generated, approved
    staff_count: int = 0
    staff_agents: int = 0
    staff_team_leads: int = 0
    staff_managers: int = 0
    operating_hours: str = ""
    location: str = ""
    pricing_model: str = "per agent per month"
    rate_agent: float = 0.0
    rate_team_lead: float = 0.0
    rate_manager: float = 0.0
    kpis: List[str] = field(default_factory=list)

    # Short Form-specific
    rates_table: List[Dict] = field(default_factory=list)

    def to_placeholders(self) -> Dict[str, str]:
        """Convert data to placeholder dictionary for template replacement."""
        return {
            "{{CLIENT_NAME}}": self.client_name,
            "{{CLIENT_LEGAL_NAME}}": self.client_legal_name or self.client_name,
            "{{CLIENT_ADDRESS}}": self.client_address,
            "{{CLIENT_CONTACT_NAME}}": self.client_contact_name,
            "{{CLIENT_CONTACT_EMAIL}}": self.client_contact_email,
            "{{CLIENT_CONTACT_PHONE}}": self.client_contact_phone,
            "{{EFFECTIVE_DATE}}": self.effective_date,
            "{{PROVIDER_NAME}}": self.provider_name,
            "{{PROVIDER_EMAIL}}": self.provider_email,
            "{{PROVIDER_PHONE}}": self.provider_phone,
            "{{PROVIDER_COMPANY}}": self.provider_company,
            "{{PROVIDER_ADDRESS}}": self.provider_address,
            # Legacy placeholders for compatibility
            "{{CATCH_NAME}}": self.provider_name,
            "{{CATCH_EMAIL}}": self.provider_email,
            "{{CATCH_PHONE}}": self.provider_phone,
            "{{CATCH_COMPANY}}": self.provider_company,
            "{{CATCH_ADDRESS}}": self.provider_address,
            "{{TERM_DURATION}}": self.term_duration,
            "{{NOTICE_PERIOD}}": self.notice_period,
            "{{JURISDICTION}}": self.jurisdiction,
            "{{PAYMENT_TERMS}}": self.payment_terms,
            "{{CURRENCY}}": self.currency,
            "{{CUSTOM_CLAUSES}}": self.custom_clauses,
            "{{PROJECT_NAME}}": self.project_name,
            "{{PROJECT_START}}": self.project_start,
            "{{PROJECT_END}}": self.project_end,
            "{{SCOPE_FORMAL}}": self.scope_formal,
            "{{STAFF_COUNT}}": str(self.staff_count),
            "{{STAFF_AGENTS}}": str(self.staff_agents),
            "{{STAFF_TEAM_LEADS}}": str(self.staff_team_leads),
            "{{STAFF_MANAGERS}}": str(self.staff_managers),
            "{{OPERATING_HOURS}}": self.operating_hours,
            "{{LOCATION}}": self.location,
            "{{PRICING_MODEL}}": self.pricing_model,
            "{{RATE_AGENT}}": f"{self.currency} {self.rate_agent:,.2f}" if self.rate_agent else "",
            "{{RATE_TEAM_LEAD}}": f"{self.currency} {self.rate_team_lead:,.2f}" if self.rate_team_lead else "",
            "{{RATE_MANAGER}}": f"{self.currency} {self.rate_manager:,.2f}" if self.rate_manager else "",
            "{{TODAY_DATE}}": date.today().strftime("%d %B %Y"),
        }


class ContractService:
    """Service for contract generation and management."""

    def __init__(self):
        self.templates_dir = CONTRACT_TEMPLATES_DIR
        self.outputs_dir = OUTPUTS_DIR / "contracts"
        self.outputs_dir.mkdir(parents=True, exist_ok=True)

        # Initialize Anthropic client if API key is available
        self.anthropic_client = None
        if ANTHROPIC_API_KEY:
            self.anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    def list_templates(self) -> List[str]:
        """List available template files."""
        if not self.templates_dir.exists():
            return []
        return [f.stem for f in self.templates_dir.glob("*.docx") if not f.name.startswith("~")]

    def load_template(self, template_name: str) -> Document:
        """Load a template document."""
        template_path = self.templates_dir / f"{template_name}.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        return Document(template_path)

    def find_placeholders(self, doc: Document) -> List[str]:
        """Find all {{PLACEHOLDER}} patterns in document."""
        placeholders = set()
        pattern = re.compile(r'\{\{[A-Z_]+\}\}')

        # Check paragraphs
        for para in doc.paragraphs:
            matches = pattern.findall(para.text)
            placeholders.update(matches)

        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = pattern.findall(para.text)
                        placeholders.update(matches)

        return sorted(list(placeholders))

    def replace_placeholders(self, doc: Document, replacements: Dict[str, str]) -> Document:
        """Replace all placeholders in document with values."""

        def replace_in_paragraph(paragraph):
            """Replace placeholders in a paragraph while preserving formatting."""
            full_text = paragraph.text
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, value or "")

            # Only modify if changes were made
            if full_text != paragraph.text:
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    for run in paragraph.runs[1:]:
                        run.text = ""

        # Process paragraphs
        for para in doc.paragraphs:
            replace_in_paragraph(para)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)

        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for para in header.paragraphs:
                        replace_in_paragraph(para)
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        replace_in_paragraph(para)

        return doc

    def generate_contract(self, contract_data: ContractData) -> Path:
        """Generate a contract from template and data."""
        # Determine template based on contract type
        template_map = {
            "MSA": "MSA_Template",
            "SOW": "SOW_Template",
            "ShortForm": "ShortForm_Template"
        }

        template_name = template_map.get(contract_data.contract_type)
        if not template_name:
            raise ValueError(f"Unknown contract type: {contract_data.contract_type}")

        # Load and populate template
        doc = self.load_template(template_name)
        replacements = contract_data.to_placeholders()
        doc = self.replace_placeholders(doc, replacements)

        # Generate filename
        safe_client = re.sub(r'[^\w\s-]', '', contract_data.client_name).strip()
        safe_client = re.sub(r'[-\s]+', '_', safe_client)
        filename = f"{safe_client}_{contract_data.contract_type}_{date.today().strftime('%Y-%m-%d')}.docx"
        output_path = self.outputs_dir / filename

        # Handle duplicate filenames
        counter = 1
        while output_path.exists():
            filename = f"{safe_client}_{contract_data.contract_type}_{date.today().strftime('%Y-%m-%d')}_{counter}.docx"
            output_path = self.outputs_dir / filename
            counter += 1

        # Save document
        doc.save(output_path)

        return output_path

    def log_contract(self, contract_data: ContractData, output_path: Path, user: str = "system") -> dict:
        """Log generated contract to JSON log file."""
        log_path = self.outputs_dir / "contract_log.json"

        # Load existing log or create new
        if log_path.exists():
            with open(log_path) as f:
                log = json.load(f)
        else:
            log = {"contracts": []}

        # Add entry
        entry = {
            "id": f"contract_{len(log['contracts']) + 1:04d}",
            "timestamp": datetime.now().isoformat(),
            "contract_type": contract_data.contract_type,
            "client_name": contract_data.client_name,
            "generated_by": user,
            "filename": output_path.name,
            "filepath": str(output_path),
            "ai_assisted_fields": ["scope_formal"] if contract_data.scope_formal else []
        }
        log["contracts"].append(entry)

        # Save log
        with open(log_path, "w") as f:
            json.dump(log, f, indent=2)

        return entry

    def get_contract_history(self, limit: int = 50) -> List[dict]:
        """Get contract generation history."""
        log_path = self.outputs_dir / "contract_log.json"
        if not log_path.exists():
            return []

        with open(log_path) as f:
            log = json.load(f)

        contracts = log.get("contracts", [])
        return list(reversed(contracts[-limit:]))

    def get_contract_by_id(self, contract_id: str) -> Optional[dict]:
        """Get a specific contract by ID."""
        log_path = self.outputs_dir / "contract_log.json"
        if not log_path.exists():
            return None

        with open(log_path) as f:
            log = json.load(f)

        for contract in log.get("contracts", []):
            if contract.get("id") == contract_id:
                return contract
        return None

    def generate_ai_scope(self, bullets: str, project_name: str) -> str:
        """Use Anthropic to generate formal scope language from bullet points."""
        if not self.anthropic_client:
            return "[AI unavailable - please write scope manually]"

        if not bullets.strip():
            return "[Please enter scope bullet points first]"

        prompt = f"""You are drafting a formal Statement of Work scope section for a professional services agreement.

Project: {project_name}

The client has provided these bullet points describing the scope:
{bullets}

Please transform these bullet points into formal contract language suitable for a Statement of Work.
- Use professional legal/business language
- Structure with lettered subsections (a), (b), (c) etc.
- Be specific and clear about deliverables
- Keep the same meaning as the bullet points
- Do not add services not mentioned in the bullets

Return ONLY the formatted scope text, no preamble or explanation."""

        try:
            message = self.anthropic_client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            return message.content[0].text
        except Exception as e:
            return f"[AI Error: {str(e)}]"

    def generate_ai_suggestions(self, contract_type: str, partial_data: dict) -> List[str]:
        """Generate AI suggestions for improving the contract."""
        if not self.anthropic_client:
            return []

        prompt = f"""You are reviewing a {contract_type} contract draft. Based on this partial data:
{json.dumps(partial_data, indent=2)}

Provide 2-3 brief suggestions (one sentence each) for improving or completing this contract.
Focus on practical business concerns. Return suggestions as a JSON array of strings."""

        try:
            message = self.anthropic_client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=300,
                messages=[{"role": "user", "content": prompt}]
            )
            # Parse suggestions from response
            response_text = message.content[0].text
            # Try to extract JSON array
            if "[" in response_text:
                start = response_text.index("[")
                end = response_text.rindex("]") + 1
                return json.loads(response_text[start:end])
            return []
        except Exception:
            return []


# Singleton instance
_contract_service: Optional[ContractService] = None


def get_contract_service() -> ContractService:
    """Get or create the contract service singleton."""
    global _contract_service
    if _contract_service is None:
        _contract_service = ContractService()
    return _contract_service
